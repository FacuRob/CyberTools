"""
Analizador de Metadatos para Archivos
Extrae información oculta de PDF, Word, Excel, imágenes y más
"""

import os
import mimetypes
from typing import Dict, List, Any, Optional
from datetime import datetime
import logging

# Imports condicionales (instalar según disponibilidad)
try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MetadataAnalyzer:
    """Analizador de metadatos para múltiples tipos de archivo"""
    
    # Tipos de archivo soportados
    SUPPORTED_TYPES = {
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
        'pdf': ['.pdf'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls'],
        'text': ['.txt', '.log', '.md']
    }
    
    def __init__(self):
        self.results = {}
        
    def analyze_file(self, file_path: str) -> Dict[str, Any]:
        """
        Analiza un archivo y extrae todos los metadatos disponibles
        
        Args:
            file_path: Ruta al archivo a analizar
            
        Returns:
            Diccionario con metadatos extraídos
        """
        if not os.path.exists(file_path):
            return {"error": "Archivo no encontrado"}
        
        # Información básica del archivo
        basic_info = self._get_basic_info(file_path)
        
        # Determinar tipo de archivo
        file_ext = os.path.splitext(file_path)[1].lower()
        file_type = self._get_file_type(file_ext)
        
        # Extraer metadatos específicos según tipo
        specific_metadata = {}
        
        if file_type == 'image' and PILLOW_AVAILABLE:
            specific_metadata = self._analyze_image(file_path)
        elif file_type == 'pdf' and PYPDF2_AVAILABLE:
            specific_metadata = self._analyze_pdf(file_path)
        elif file_type == 'word' and DOCX_AVAILABLE:
            specific_metadata = self._analyze_word(file_path)
        elif file_type == 'excel' and OPENPYXL_AVAILABLE:
            specific_metadata = self._analyze_excel(file_path)
        elif file_type == 'text':
            specific_metadata = self._analyze_text(file_path)
        
        # Combinar resultados
        return {
            "status": "success",
            "file_info": basic_info,
            "file_type": file_type,
            "metadata": specific_metadata,
            "warnings": self._check_security_risks(specific_metadata),
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def _get_basic_info(self, file_path: str) -> Dict[str, Any]:
        """Extrae información básica del archivo"""
        stat = os.stat(file_path)
        
        return {
            "filename": os.path.basename(file_path),
            "size": self._format_size(stat.st_size),
            "size_bytes": stat.st_size,
            "extension": os.path.splitext(file_path)[1].lower(),
            "mime_type": mimetypes.guess_type(file_path)[0] or "unknown",
            "created": datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
            "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            "accessed": datetime.fromtimestamp(stat.st_atime).strftime("%Y-%m-%d %H:%M:%S")
        }
    
    def _get_file_type(self, extension: str) -> str:
        """Determina el tipo de archivo por extensión"""
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if extension in extensions:
                return file_type
        return "unknown"
    
    def _analyze_image(self, file_path: str) -> Dict[str, Any]:
        """Extrae metadatos EXIF de imágenes"""
        try:
            image = Image.open(file_path)
            
            # Información básica de la imagen
            info = {
                "dimensions": f"{image.width}x{image.height}",
                "format": image.format or "Unknown",
                "mode": image.mode,
            }
            
            # Extraer EXIF (método compatible con más versiones de Pillow)
            exif_data = {}
            try:
                # Método 1: _getexif (antiguo)
                if hasattr(image, '_getexif'):
                    exif_raw = image._getexif()
                    if exif_raw:
                        for tag_id, value in exif_raw.items():
                            tag = TAGS.get(tag_id, str(tag_id))
                            try:
                                if isinstance(value, bytes):
                                    value = value.decode('utf-8', errors='ignore')
                                exif_data[tag] = str(value)
                            except:
                                pass
                
                # Método 2: getexif (moderno)
                elif hasattr(image, 'getexif'):
                    exif_raw = image.getexif()
                    if exif_raw:
                        for tag_id, value in exif_raw.items():
                            tag = TAGS.get(tag_id, str(tag_id))
                            try:
                                if isinstance(value, bytes):
                                    value = value.decode('utf-8', errors='ignore')
                                exif_data[tag] = str(value)
                            except:
                                pass
                
                # Método 3: info (fallback)
                if not exif_data and hasattr(image, 'info'):
                    for key, value in image.info.items():
                        try:
                            if isinstance(value, (str, int, float)):
                                exif_data[key] = str(value)
                        except:
                            pass
                            
            except Exception as exif_error:
                logger.warning(f"Error extracting EXIF: {exif_error}")
            
            # Determinar si hay EXIF
            if exif_data:
                info["exif"] = exif_data
                info["exif_count"] = len(exif_data)
            else:
                info["exif"] = None
                info["exif_note"] = "No EXIF data found (puede haber sido eliminado por WhatsApp/redes sociales)"
            
            # Datos potencialmente sensibles
            sensitive = {}
            if exif_data:
                sensitive_tags = ['GPSInfo', 'Make', 'Model', 'Software', 
                                 'DateTime', 'Artist', 'Copyright', 'DateTimeOriginal']
                for tag in sensitive_tags:
                    if tag in exif_data:
                        sensitive[tag] = exif_data[tag]
            
            info["sensitive_data"] = sensitive if sensitive else None
            
            # Información adicional de la imagen
            if hasattr(image, 'info') and image.info:
                additional = {}
                safe_keys = ['dpi', 'compression', 'quality', 'jfif', 'jfif_version']
                for key in safe_keys:
                    if key in image.info:
                        try:
                            additional[key] = str(image.info[key])
                        except:
                            pass
                if additional:
                    info["additional_info"] = additional
            
            return info
            
        except Exception as e:
            logger.error(f"Error analyzing image: {e}", exc_info=True)
            return {"error": f"No se pudo analizar la imagen: {str(e)}"}
    
    def _analyze_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extrae metadatos de archivos PDF"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                info = {
                    "pages": len(reader.pages),
                    "encrypted": reader.is_encrypted
                }
                
                # Metadatos del documento
                if reader.metadata:
                    metadata = {}
                    for key, value in reader.metadata.items():
                        # Limpiar claves (quitar el prefijo /)
                        clean_key = key.replace('/', '') if key.startswith('/') else key
                        metadata[clean_key] = str(value)
                    
                    info["document_info"] = metadata
                    
                    # Datos sensibles comunes
                    sensitive = {}
                    sensitive_keys = ['Author', 'Creator', 'Producer', 'Title', 
                                     'Subject', 'Keywords', 'CreationDate', 'ModDate']
                    for key in sensitive_keys:
                        if key in metadata:
                            sensitive[key] = metadata[key]
                    
                    info["sensitive_data"] = sensitive if sensitive else None
                
                return info
                
        except Exception as e:
            logger.error(f"Error analyzing PDF: {e}")
            return {"error": str(e)}
    
    def _analyze_word(self, file_path: str) -> Dict[str, Any]:
        """Extrae metadatos de documentos Word (.docx)"""
        try:
            doc = Document(file_path)
            
            info = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            # Propiedades del documento
            core_props = doc.core_properties
            
            properties = {
                "author": core_props.author or "Unknown",
                "title": core_props.title or "Untitled",
                "subject": core_props.subject or "None",
                "keywords": core_props.keywords or "None",
                "created": str(core_props.created) if core_props.created else "Unknown",
                "modified": str(core_props.modified) if core_props.modified else "Unknown",
                "last_modified_by": core_props.last_modified_by or "Unknown",
                "revision": core_props.revision or "Unknown"
            }
            
            info["document_properties"] = properties
            
            # Datos sensibles
            sensitive = {k: v for k, v in properties.items() 
                        if v and v not in ["Unknown", "Untitled", "None"]}
            info["sensitive_data"] = sensitive if sensitive else None
            
            return info
            
        except Exception as e:
            logger.error(f"Error analyzing Word document: {e}")
            return {"error": str(e)}
    
    def _analyze_excel(self, file_path: str) -> Dict[str, Any]:
        """Extrae metadatos de archivos Excel (.xlsx)"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            info = {
                "sheets": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "active_sheet": workbook.active.title
            }
            
            # Propiedades del documento
            props = workbook.properties
            
            properties = {
                "creator": props.creator or "Unknown",
                "title": props.title or "Untitled",
                "subject": props.subject or "None",
                "description": props.description or "None",
                "keywords": props.keywords or "None",
                "created": str(props.created) if props.created else "Unknown",
                "modified": str(props.modified) if props.modified else "Unknown",
                "last_modified_by": props.lastModifiedBy or "Unknown"
            }
            
            info["document_properties"] = properties
            
            # Datos sensibles
            sensitive = {k: v for k, v in properties.items() 
                        if v and v not in ["Unknown", "Untitled", "None"]}
            info["sensitive_data"] = sensitive if sensitive else None
            
            workbook.close()
            
            return info
            
        except Exception as e:
            logger.error(f"Error analyzing Excel file: {e}")
            return {"error": str(e)}
    
    def _analyze_text(self, file_path: str) -> Dict[str, Any]:
        """Analiza archivos de texto plano"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                lines = content.split('\n')
            
            return {
                "lines": len(lines),
                "characters": len(content),
                "words": len(content.split()),
                "encoding": "UTF-8 (detected)"
            }
            
        except Exception as e:
            logger.error(f"Error analyzing text file: {e}")
            return {"error": str(e)}
    
    def _check_security_risks(self, metadata: Dict[str, Any]) -> List[str]:
        """Identifica riesgos de seguridad en los metadatos"""
        warnings = []
        
        if not metadata or "error" in metadata:
            return warnings
        
        # Verificar datos sensibles
        if "sensitive_data" in metadata and metadata["sensitive_data"]:
            sensitive = metadata["sensitive_data"]
            
            if any(key in sensitive for key in ['Author', 'Creator', 'author', 'creator']):
                warnings.append("⚠️ Contiene información del autor")
            
            if any(key in sensitive for key in ['GPSInfo', 'Location']):
                warnings.append("🚨 CRÍTICO: Contiene datos de geolocalización")
            
            if any(key in sensitive for key in ['Make', 'Model', 'Software', 'Producer']):
                warnings.append("⚠️ Revela software/dispositivo usado")
            
            if 'Copyright' in sensitive:
                warnings.append("ℹ️ Contiene información de copyright")
        
        # Verificar EXIF en imágenes
        if "exif" in metadata and isinstance(metadata["exif"], dict):
            if metadata["exif"]:
                warnings.append("⚠️ Contiene metadatos EXIF (potencialmente sensibles)")
        
        return warnings if warnings else ["✅ No se detectaron riesgos evidentes"]
    
    def _format_size(self, bytes_size: int) -> str:
        """Formatea el tamaño del archivo en formato legible"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if bytes_size < 1024.0:
                return f"{bytes_size:.2f} {unit}"
            bytes_size /= 1024.0
        return f"{bytes_size:.2f} TB"


def analyze_metadata(file_path: str) -> Dict[str, Any]:
    """
    Función principal para analizar metadatos de un archivo
    
    Args:
        file_path: Ruta al archivo
        
    Returns:
        Diccionario con resultados del análisis
    """
    analyzer = MetadataAnalyzer()
    return analyzer.analyze_file(file_path)


# Test básico
if __name__ == "__main__":
    print("=== Analizador de Metadatos ===")
    print(f"PIL/Pillow: {'✓' if PILLOW_AVAILABLE else '✗'}")
    print(f"PyPDF2: {'✓' if PYPDF2_AVAILABLE else '✗'}")
    print(f"python-docx: {'✓' if DOCX_AVAILABLE else '✗'}")
    print(f"openpyxl: {'✓' if OPENPYXL_AVAILABLE else '✗'}")